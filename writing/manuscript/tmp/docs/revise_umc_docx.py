# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/ujunbin/project/umc/Writing/manuscript")
INPUT = ROOT / "UMC_report_draft_kr.docx"
WORK = ROOT / "output/doc/UMC_report_draft_kr_revised_work.docx"
STRIPPED = ROOT / "output/doc/UMC_report_draft_kr_revised_stripped.docx"
OUTPUT = ROOT / "output/doc/UMC_report_draft_kr_revised.docx"
FIG_DIR = ROOT / "output/doc/figures"
REPORT_FIG_DIR = Path("/Users/ujunbin/project/umc/Analysis/Part 1/output/figures/report")
TABLE_IMAGE_DIR = ROOT / "output/doc/table_images"
SKILL_DIR = Path(
    "/Users/ujunbin/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.423.10653/skills/documents"
)


def set_paragraph_text(paragraph, text: str, *, bold: bool = False) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def replace_exact(doc: Document, old: str, new: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == old.strip():
            set_paragraph_text(paragraph, new)
            return True
    return False


def replace_startswith(doc: Document, prefix: str, new: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_paragraph_text(paragraph, new)
            return True
    return False


def blank_startswith(doc: Document, prefix: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_paragraph_text(paragraph, "")
            return True
    return False


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    set_paragraph_text(new_paragraph, text)
    new_paragraph.style = paragraph.style
    return new_paragraph


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: float = 8.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def set_table_borders(table, color: str = "D6DCE5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size, index=1 if bold and candidate.endswith(".ttc") else 0)
    return ImageFont.load_default()


def text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    if not text:
        return 0
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0]


def wrap_one_line(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    words = text.split(" ")
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_width(draw, candidate, font) <= max_width:
            current = candidate
            continue
        if current:
            lines.append(current)
            current = ""
        if text_width(draw, word, font) <= max_width:
            current = word
            continue
        chunk = ""
        for ch in word:
            candidate = chunk + ch
            if text_width(draw, candidate, font) <= max_width:
                chunk = candidate
            else:
                if chunk:
                    lines.append(chunk)
                chunk = ch
        current = chunk
    if current:
        lines.append(current)
    return lines or [""]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    text = str(text).replace("\r", "").strip()
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap_one_line(draw, part.strip(), font, max_width))
    return lines or [""]


def extract_table_rows(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([
            cell.text.strip()
            .replace("\r", "")
            .replace("−", "-")
            .replace("–", "-")
            for cell in row.cells
        ])
    return rows


def clean_repeated_section_rows(rows: list[list[str]]) -> list[list[str]]:
    cleaned: list[list[str]] = []
    for row in rows:
        nonempty = [cell.strip() for cell in row if cell.strip()]
        if nonempty and len(set(nonempty)) == 1 and ("Level" in nonempty[0] or "Variable" not in nonempty[0]):
            cleaned.append([nonempty[0]] + [""] * (len(row) - 1))
        else:
            cleaned.append(row)
    return cleaned


def draw_table_image(
    rows: list[list[str]],
    col_widths: list[int],
    out_path: Path,
    *,
    font_size: int = 26,
    header_rows: int = 1,
    title: str | None = None,
) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    margin_x = 36
    margin_y = 32
    pad_x = 14
    pad_y = 10
    body_font = load_font(font_size)
    header_font = load_font(font_size, bold=True)
    title_font = load_font(font_size + 6, bold=True)

    probe = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(probe)
    line_h = int(font_size * 1.32)
    wrapped_rows: list[list[list[str]]] = []
    heights: list[int] = []
    n_cols = len(col_widths)
    for r_idx, row in enumerate(rows):
        row = (row + [""] * n_cols)[:n_cols]
        font = header_font if r_idx < header_rows else body_font
        wrapped = [wrap_text(draw, cell, font, col_widths[c] - pad_x * 2) for c, cell in enumerate(row)]
        wrapped_rows.append(wrapped)
        heights.append(max(44, max(len(lines) for lines in wrapped) * line_h + pad_y * 2))

    title_h = 0
    if title:
        title_h = line_h + 18
    width = sum(col_widths) + margin_x * 2
    height = sum(heights) + margin_y * 2 + title_h
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    y = margin_y
    if title:
        draw.text((margin_x, y), title, fill=(25, 39, 70), font=title_font)
        y += title_h

    x_positions = [margin_x]
    for width_i in col_widths[:-1]:
        x_positions.append(x_positions[-1] + width_i)

    border = (198, 207, 219)
    header_fill = (232, 238, 247)
    section_fill = (244, 247, 251)
    alt_fill = (250, 252, 255)
    text_color = (32, 38, 48)
    for r_idx, row_lines in enumerate(wrapped_rows):
        row_h = heights[r_idx]
        is_section = r_idx >= header_rows and row_lines[0][0].startswith("Level ")
        fill = header_fill if r_idx < header_rows else section_fill if is_section else alt_fill if r_idx % 2 == 0 else (255, 255, 255)
        draw.rectangle([margin_x, y, margin_x + sum(col_widths), y + row_h], fill=fill, outline=border)
        x = margin_x
        for c_idx, lines in enumerate(row_lines):
            draw.rectangle([x, y, x + col_widths[c_idx], y + row_h], outline=border)
            font = header_font if r_idx < header_rows or is_section else body_font
            tx = x + pad_x
            ty = y + pad_y
            for line in lines:
                draw.text((tx, ty), line, fill=text_color, font=font)
                ty += line_h
            x += col_widths[c_idx]
        y += row_h
    image.save(out_path)


def insert_paragraph_after_xml(xml_element, parent, *, text: str | None = None, image: Path | None = None, width_cm: float = 16.5):
    new_p = OxmlElement("w:p")
    xml_element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text is not None:
        set_paragraph_text(paragraph, text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
    if image is not None:
        paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
        paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def insert_blocks_after_table(table, blocks: list[dict]) -> None:
    anchor = table._element
    parent = table._parent
    for block in blocks:
        paragraph = insert_paragraph_after_xml(anchor, parent, **block)
        anchor = paragraph._p
    table._element.getparent().remove(table._element)


def replace_problematic_tables_with_images(doc: Document) -> None:
    TABLE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    tables = list(doc.tables)

    table_plan: dict[int, list[dict]] = {}

    # Table 1: long indicator table, split across two readable images.
    t1_rows = extract_table_rows(tables[1])
    t1_cols = [300, 390, 470, 440, 680, 100]
    t1a = TABLE_IMAGE_DIR / "table1_indicators_part1.png"
    t1b = TABLE_IMAGE_DIR / "table1_indicators_part2.png"
    draw_table_image([t1_rows[0]] + t1_rows[1:9], t1_cols, t1a, font_size=22)
    draw_table_image([t1_rows[0]] + t1_rows[9:], t1_cols, t1b, font_size=22)
    table_plan[1] = [{"image": t1a}, {"image": t1b}]

    # Original table wrappers for paired figures are fragile in the artifact renderer.
    table_plan[2] = [
        {"text": "Figure 2. UMC Composite Scores by District"},
        {"image": REPORT_FIG_DIR / "fig_report_bar_umc.png", "width_cm": 12.5},
        {"text": "Figure 3. Heatmap of UMC Dimension Scores"},
        {"image": REPORT_FIG_DIR / "fig_report_heatmap.png", "width_cm": 12.5},
    ]
    table_plan[3] = [{"image": FIG_DIR / "fig_lisa_dimensions_en.png", "width_cm": 16.0}]

    # Data tables.
    t4_rows = extract_table_rows(tables[4])
    t4_cols = [390, 720, 130, 160, 420]
    t4a = TABLE_IMAGE_DIR / "table2_variable_summary_part1.png"
    t4b = TABLE_IMAGE_DIR / "table2_variable_summary_part2.png"
    draw_table_image(t4_rows[:10], t4_cols, t4a, font_size=23)
    draw_table_image([t4_rows[0]] + t4_rows[10:], t4_cols, t4b, font_size=23)
    table_plan[4] = [{"image": t4a}, {"image": t4b}]

    t5_rows = extract_table_rows(tables[5])
    t5_cols = [560, 230, 230, 230, 230, 230]
    t5_img = TABLE_IMAGE_DIR / "table3_descriptive_statistics.png"
    draw_table_image(t5_rows, t5_cols, t5_img, font_size=24)
    table_plan[5] = [{"image": t5_img}]

    t6_rows = clean_repeated_section_rows(extract_table_rows(tables[6]))
    t6_cols = [520, 150, 330, 330, 330]
    t6a = TABLE_IMAGE_DIR / "table4_hlm_part1.png"
    t6b = TABLE_IMAGE_DIR / "table4_hlm_part2.png"
    draw_table_image(t6_rows[:18], t6_cols, t6a, font_size=22)
    draw_table_image([t6_rows[0]] + t6_rows[18:], t6_cols, t6b, font_size=22)
    table_plan[6] = [{"image": t6a}, {"image": t6b}]

    t7_rows = extract_table_rows(tables[7])
    t7_img = TABLE_IMAGE_DIR / "table5_units_analysis.png"
    draw_table_image(t7_rows, [300, 360, 560, 720], t7_img, font_size=24)
    table_plan[7] = [{"image": t7_img}]

    t8_rows = extract_table_rows(tables[8])
    t8_img = TABLE_IMAGE_DIR / "table_posterior_shift_rules.png"
    draw_table_image(t8_rows, [410, 430, 760], t8_img, font_size=24)
    table_plan[8] = [{"image": t8_img}]

    t9_rows = extract_table_rows(tables[9])
    t9_cols = [250, 280, 200, 200, 230, 230]
    t9a = TABLE_IMAGE_DIR / "table_text_classification_part1.png"
    t9b = TABLE_IMAGE_DIR / "table_text_classification_part2.png"
    draw_table_image([t9_rows[0]] + t9_rows[1:14], t9_cols, t9a, font_size=24)
    draw_table_image([t9_rows[0]] + t9_rows[14:], t9_cols, t9b, font_size=24)
    table_plan[9] = [{"text": "Table 6. Text Classification Results by District"}, {"image": t9a}, {"image": t9b}]

    t10_rows = extract_table_rows(tables[10])
    t10_img = TABLE_IMAGE_DIR / "table6_hypotheses.png"
    draw_table_image(t10_rows, [510, 430, 430, 220], t10_img, font_size=24)
    table_plan[10] = [{"image": t10_img}]

    for idx in sorted(table_plan.keys(), reverse=True):
        insert_blocks_after_table(tables[idx], table_plan[idx])


def resize_small_formula_images(doc: Document) -> None:
    for shape in doc.inline_shapes:
        if shape.width < Cm(7) and shape.height < Cm(2):
            shape.width = Cm(10)


def remove_trailing_empty_paragraphs(doc: Document) -> None:
    body = doc._body._element
    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            text = "".join(child.xpath(".//w:t/text()"))
            has_drawing = bool(child.xpath(".//w:drawing"))
            if not text.strip() and not has_drawing:
                body.remove(child)
                continue
        break


def rebuild_indicator_table(doc: Document) -> None:
    table = doc.tables[1]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    headers = [
        "Dimension",
        "Indicator",
        "Unit / Raw scale",
        "Source / year",
        "Construction / weighting",
        "Direction",
    ]
    while len(table.columns) < len(headers):
        table.add_column(Cm(2.0))
    # If the original table has fewer columns, add_column handles it; if more, keep only first six visible.
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, center=True, size=7.6)
        shade_cell(table.rows[0].cells[i], "E8EEF7")

    rows = [
        ("Connectivity", "4G station density", "Stations per 1,000 daytime living population", "Spectrum Resource Management System; Seoul daytime living population", "Unique 4G station coordinates divided by daytime living population", "+"),
        ("Connectivity", "5G station density", "Stations per 1,000 daytime living population", "Spectrum Resource Management System; Seoul daytime living population", "Unique 5G station coordinates divided by daytime living population", "+"),
        ("Connectivity", "Download speed (CQ)", "Mbps", "NIA Communication Quality (CQ), 2023-2024", "District mean of measurement points", "+"),
        ("Available for Use", "Mobile data usage", "GB per household-month", "SKT Telecom Data, 2023-2024", "Population-weighted district monthly average, then annual mean", "+"),
        ("Available for Use", "Online service use days", "Days per month", "SKT Telecom Data, 2023-2024", "Population-weighted use days for finance, shopping, video, and delivery services", "+"),
        ("Available for Use", "Public WiFi AP density", "APs per 1,000 residents", "Local Administrative WiFi License Data", "Registered public WiFi access points divided by resident population", "+"),
        ("Affordability", "Average household income", "KRW per month", "Seoul Commercial Analysis Service", "District monthly household income from commercial-area administrative statistics", "+"),
        ("Affordability", "Recent fee delinquency", "% delinquent in last three months", "SKT Telecom Data", "Population-weighted district average; reverse-normalized", "-"),
        ("Devices", "Core device ownership at home", "Count, 0-3", "Seoul Digital Competency Survey, 2023", "Weighted district mean using survey weight WT; desktop, laptop, tablet", "+"),
        ("Devices", "Core device use", "Count, 0-4", "Seoul Digital Competency Survey, 2023", "Weighted district mean using WT; desktop, laptop, smartphone, tablet", "+"),
        ("Digital Skills", "Basic skills", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of Q4 items divided by maximum possible score of 24; weighted mean using WT", "+"),
        ("Digital Skills", "Applied skills", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of Q5B items divided by maximum possible score of 40; weighted mean using WT", "+"),
        ("Digital Skills", "Advanced skills", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of Q6B items divided by maximum possible score of 28; weighted mean using WT", "+"),
        ("Digital Skills", "Cyber activity", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of Q7 items divided by maximum possible score of 20; weighted mean using WT", "+"),
        ("Safety", "Security behavior", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of seven behavior items divided by maximum possible score of 28; weighted mean using WT", "+"),
        ("Safety", "Security awareness", "Ratio, 0-1", "Seoul Digital Competency Survey, 2023", "Sum of three awareness items divided by maximum possible score of 12; weighted mean using WT", "+"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, center=(i == 5), size=7.1)
            if i == 0:
                shade_cell(cells[i], "F5F7FA")

    widths = [2.25, 2.7, 3.3, 3.05, 4.95, 1.0]
    for row in table.rows:
        for i, width in enumerate(widths):
            set_cell_width(row.cells[i], width)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    set_table_borders(table)


def update_variable_table(doc: Document) -> None:
    table = doc.tables[4]
    while len(table.rows) > 13:
        table._tbl.remove(table.rows[-1]._tr)
    rows = [
        ("connectivity", "Connectivity dimension score", "L2", "0-1", "UMC Part 1"),
        ("available_for_use", "Available for Use dimension score", "L2", "0-1", "UMC Part 1"),
        ("affordability", "Affordability dimension score", "L2", "0-1", "UMC Part 1"),
        ("district_income", "District average household income", "L2", "0-1", "Seoul Commercial Analysis Service"),
        ("aging_rate", "Share of residents aged 65+", "L2", "0-1", "Population Statistics"),
    ]
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, center=(i in {2, 3}), size=8.0)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, [3.3, 5.5, 1.6, 2.0, 3.9][i])
    for cell in table.rows[0].cells:
        shade_cell(cell, "E8EEF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    set_table_borders(table)


def polish_styles(doc: Document) -> None:
    for style_name in ("Normal", "normal"):
        try:
            style = doc.styles[style_name]
            style.font.name = "맑은 고딕"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            style.font.size = Pt(10)
        except Exception:
            pass
    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        if paragraph.style.name.lower().startswith("heading"):
            pf.space_before = Pt(10)
            pf.space_after = Pt(5)
        else:
            pf.space_after = Pt(6)
            pf.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def replace_text(doc: Document) -> None:
    replace_startswith(
        doc,
        "그 다음, 개인 특성과 자치구 수준의 UMC 인프라 특성을",
        "그 다음, 개인 특성과 자치구 수준의 UMC 인프라 특성을 함께 투입한 2수준 위계적 선형 모형(HLM)을 통해, 디지털 활용의 차이가 개인 요인 및 지역 요인과 어떻게 관련되는지 검토한다. 본 분석은 관찰자료에 기반한 회귀분석이므로 인과 효과를 식별하기보다, 개인 특성과 지역 조건의 상대적 설명력과 취약집단에서의 차별적 연관을 추정하는 데 목적을 둔다. HLM은 개인이 지역에 내포된 위계적 자료 구조를 고려하는 표준적 방법이며, 투입 변수를 해당 맥락의 인구 특성과 인프라 조건에 맞게 교체하면 동일한 분석 틀을 다른 도시에도 적용할 수 있다. 변수 선정, 모형 설정, 교차수준 상호작용의 구체적 설계는 3.2절에서 상술한다.",
    )
    replace_startswith(
        doc,
        "마지막으로, 비정형 텍스트를 통해 앞선 두 분석이 포착하지 못하는",
        "마지막으로, 비정형 텍스트를 통해 앞선 두 분석이 포착하지 못하는 개별 현상을 복원하고, UMC 개념에 적합한 상위 구조로 추상화한다. 개별 사례를 지역 수준의 UMC 개념과 곧바로 동일시하지 않고 집계와 대조 과정을 거치며, 앞선 관측 결과가 허용하는 범위 안에서 가능한 메커니즘을 탐색한다. 따라서 텍스트 분석은 인과 식별을 대체하는 절차가 아니라, 회귀 기반 연관 분석만으로 설명하기 어려운 생활세계의 맥락을 보완하는 단계이다.",
    )
    replace_startswith(
        doc,
        "구체적으로, 3.1절의 UMC 실측값을 사전 분포로",
        "구체적으로, 3.1절의 UMC 실측값을 사전 분포로, 텍스트 분류 결과를 우도로 결합하여 지역별 UMC 표출 패턴을 베이지안으로 갱신한다. 이어서, 3.2절의 HLM 결과가 허용하는 범위 내에서 구조와 경험 사이의 가능한 메커니즘을 가시화한다. 이때 귀추적, 전향적, 경로적 추론의 세 가지 독립된 추론층을 설계하여, 각 추론 에이전트가 서로의 출력을 참조하지 않은 상태에서 가설을 생성한 뒤 상위 판정 에이전트가 이를 종합한다. 추론층을 분리함으로써 단일 추론 경로의 편향을 억제하고, 가설 간 수렴과 발산을 관찰할 수 있다. 프롬프트 설계, 오분류 억제 전략, 베이지안 업데이트 및 추론 절차의 구체적 내용은 3.3절에서 상술한다.",
    )
    replace_startswith(
        doc,
        "본 장의 세 분석은 순차적으로 연결된다.",
        "본 장의 세 분석은 순차적으로 연결된다. 먼저 서울 자치구별 UMC 지수를 구축하여 지역 격차를 측정하고(3.1), 이 결과를 지역 수준 변수로 활용하여 개인 및 지역 특성이 개인의 디지털 활용과 어떻게 관련되는지 분석한다(3.2). 마지막으로 디지털 플랫폼 텍스트를 통해 양적 분석이 포착하지 못하는 구체적 경험과 표출 양상을 검토한다. 이 흐름은 측정 -> 회귀 기반 연관 분석 -> 생활세계 맥락 보완으로 구성되며, 3.4절에서 세 분석의 결과를 종합한다.",
    )
    replace_startswith(
        doc,
        "측정 지표 선정 과정에서는 개념적 정의와의 부합성",
        "측정 지표 선정 과정에서는 개념적 정의와의 부합성, 서울의 디지털 기술 환경, 데이터 가용성을 함께 검토하였다. 분석에는 2023년과 2024년 자료를 활용하되, Devices, Digital Skills, Safety는 서울시민 디지털역량실태조사의 제공 시점상 2023년 값을 두 연도에 공통 적용하였다. 따라서 본 지수는 서울 25개 자치구 간 상대적 위치를 비교하기 위한 도시 내부 지표이며, 국가 간 또는 시점 간 절대 수준을 직접 비교하는 지표로 해석하지 않는다.",
    )
    replace_startswith(
        doc,
        "Connectivity의 경우, 한국전파누리에서 제공하는",
        "Connectivity는 한국전파누리/Spectrum Resource Management System의 이동통신 무선국 좌표 자료와 NIA 통신서비스 품질평가(CQ) 자료를 결합해 산출하였다. 4G 및 5G 기지국 지표는 중복 좌표를 제거한 고유 기지국 수를 주간 생활인구 1,000명당 밀도로 환산했으며, 다운로드 속도는 자치구별 측정 지점의 평균 Mbps를 사용하였다.",
    )
    replace_startswith(
        doc,
        "Available for Use는 디지털 서비스의 실질적 이용 가능성을",
        "Available for Use는 디지털 서비스의 실질적 이용 가능성을 측정하는 차원이다. SKT 통신정보의 월별 모바일 데이터 사용량과 온라인 서비스 이용일수, 그리고 자치구 공공 WiFi AP 등록 현황을 사용하였다. SKT 지표는 행정동·성·연령대별 값을 인구가중 평균하여 자치구 월별 값으로 집계한 뒤 연평균을 산출했고, 공공 WiFi는 전체 인구 1,000명당 AP 수로 환산하였다.",
    )
    replace_startswith(
        doc,
        "Affordability는 ITU ICT Development Index",
        "Affordability는 ITU ICT Development Index의 구매력(capacity)과 부담(burden) 구분을 따라 구성하였다. 구매력은 서울시가 제공하는 서울시 상권분석서비스(상권·행정자료 기반 지역 상권 통계)의 자치구별 월평균 가구소득을 사용하였다. 부담은 SKT 통신정보의 최근 3개월 이내 요금 연체 비율을 사용하고 역방향 지표로 처리하였다. 이동통신 요금제 가격 자체는 전국 단위로 책정되어 자치구별로 크게 달라지지 않으므로, 본 분석에서 자치구별 Affordability 차이는 주로 소득 수준과 연체율의 차이를 반영한다.",
    )
    replace_startswith(
        doc,
        "Digital Skills는 기초, 응용, 고급",
        "Digital Skills는 기초, 응용, 고급, 사이버활동의 4개 하위 영역으로 구성된다. 각 영역은 문항 합산점수를 해당 영역의 최대가능점수(기초 24점, 응용 40점, 고급 28점, 사이버활동 20점)로 나누어 0-1 원점수 비율로 변환하였다. 이후 서울시민 디지털역량실태조사의 조사 가중치(WT)를 적용해 자치구별 가중평균을 산출하였다.",
    )
    replace_startswith(
        doc,
        "Safety는 디지털 보안 행동과 보안 인식의",
        "Safety는 디지털 보안 행동과 보안 인식의 2개 하위 지표로 구성된다. 보안 행동은 7개 문항 합산점수를 28점으로, 보안 인식은 3개 문항 합산점수를 12점으로 나누어 0-1 원점수 비율로 변환한 뒤 조사 가중치(WT)를 적용해 자치구별 가중평균을 산출하였다.",
    )
    replace_startswith(
        doc,
        "모든 원시 지표는 25개 자치구 내에서 Min-Max",
        "모든 원시 지표는 연도별로 서울 25개 자치구의 관측 최솟값과 최댓값을 기준으로 [0, 1] 구간에 Min-Max 정규화하였다. 값이 클수록 좋은 정방향 지표는 x* = (x - min(x)) / (max(x) - min(x))로 산출하고, 값이 클수록 불리한 역방향 지표(요금 연체율)는 x* = (max(x) - x) / (max(x) - min(x))로 변환하였다. 이 방식은 서울 내부의 상대적 격차를 선명하게 보여주지만, 이론적 절대 기준을 사용한 것은 아니므로 다른 도시 또는 국가와의 직접 비교에는 주의가 필요하다.",
    )
    replace_startswith(
        doc,
        "이때, 값이 클수록 나쁜 결과를 의미하는 역방향 지표는",
        "각 차원 점수는 해당 차원에 속한 정규화 지표의 동일가중 산술평균이며, 종합 UMC 지수는 6개 차원 점수의 동일가중 평균이다. 동일가중을 적용한 이유는 서울 맥락에서 6개 차원의 상대적 중요도에 대한 합의된 외부 기준이 아직 없고, 해커톤 제출물의 목적상 연구자의 임의 가중을 최소화하는 편이 적절하기 때문이다. 산출된 종합지수는 0(최저)에서 1(최고) 범위를 가지며, 서울 25개 자치구 전체에 대해 2023년과 2024년 각각 독립적으로 산출한다.",
    )
    blank_startswith(doc, "결론적으로, Figure 4에 제시된 바와 같이 공간적으로 밀집한 종합지수 하위 4개")
    replace_startswith(
        doc,
        "종합지수의 공간적 분포를 살펴보면",
        "종합지수의 공간적 분포를 살펴보면(Figure 4), UMC 지수가 높게 나온 자치구는 두 구역에 집중되어 있었다. 이는 각각 서초-영등포-마포로 이어지는 한강과 도심권 지역과, 용산-종로-중구 중심의 전통적 도심부이다. 반면, 북부 외곽(강북구, 도봉구, 노원구)과 남서부(구로구, 금천구)는 도시 평균 이하의 점수를 보였다. 이러한 공간 분포는 서울의 사회경제적인 자원 분포와 연관되어 있었다. 즉, 디지털 격차가 도시의 구조적 공간 불평등과 밀접히 연관되어 있을 가능성을 시사하였다.",
    )
    replace_startswith(
        doc,
        "결론적으로, Figure 4에 제시된 바와 같이 종합지수 하위 5개",
        "결론적으로, Figure 4에 제시된 바와 같이 종합지수 하위 5개 자치구(중랑구, 도봉구, 강북구, 구로구, 노원구)를 본 프로젝트에서는 잠정적으로 디지털 사막(Digital Desert)으로 분류한다. 이들 자치구는 공통적으로 Devices와 Safety 차원에서 뚜렷한 결핍을 보이며, Affordability 역시 서울시 평균에 미치지 못했다. 한편, Connectivity 차원에서는 일부 자치구(중랑구 0.630, 강북구 0.562)가 서울시 평균을 상회하여, 물리적 네트워크 인프라 자체의 부재보다는 이를 활용하기 위한 인적·경제적 역량의 부족이 핵심 격차 요인일 가능성을 보여준다. 이러한 발견은 3.2절의 다층 분석에서 개인 수준 변수와 자치구 수준 변수의 상대적 기여를 분리 및 검토함으로써 보다 체계적으로 다룬다. 또한 일부 UMC 차원에서는 지리적으로 인접한 지역이 서로 유사한 패턴을 보였다(Figure 5). 공간적으로 군집한 자치구가 유사한 취약성을 보인다는 점은 디지털 사막 해소를 위한 권역 단위 거버넌스 체계 구축의 필요성을 시사한다.",
    )
    blank_startswith(doc, "또한, UMC의 일부 차원에서 지리적으로 인접한 지역이 서로 유사한 패턴을 보였다.")
    replace_startswith(
        doc,
        "두 번째 분석에서 활용한 자료인 서울서베이는",
        "두 번째 분석에서 활용한 자료인 서울서베이는 2003년부터 서울시가 매년 실시해 온 도시정책 기초조사이다. 자료 수집은 매년 8-9월에 비대면 조사와 방문면접을 병행해 이루어지며, 본 연구는 2023년과 2024년 개인 응답을 결합하여 분석하였다.",
    )
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("두 번째 분석에서 활용한 자료인 서울서베이는"):
            insert_after(
                paragraph,
                "표본 구성과 결측 처리. 분석 파일은 서울서베이 2023년 5,000명과 2024년 5,000명을 결합한 총 10,000명으로 구성되었다. 종속변수와 주요 개인 수준 예측변수(연령, 성별, 교육, 소득, 직업, 장애 여부, 가구 구성, 주거유형) 및 자치구 수준 UMC 변수에는 결측이 없었으므로 추가 대체(imputation)는 수행하지 않았다. 분석은 최종 결합 파일 기준의 완전 사례로 추정하였다. 공개 분석 파일에 적용 가능한 설계가중치가 포함되어 있지 않아, 결과는 설계가중 추정치가 아니라 자치구와 주요 인구사회학적 특성을 조건화한 모형 기반 연관으로 해석한다.",
            )
            break
    replace_startswith(
        doc,
        "지역 수준(Level 2) 변수는 3.1절에서 구축한 UMC",
        "지역 수준(Level 2) 변수는 3.1절에서 구축한 UMC 차원 중 Connectivity, Available for Use, Affordability를 주 분석에 사용하였다. Devices, Digital Skills, Safety는 서울시민 디지털역량실태조사의 개인 응답을 자치구 단위로 집계한 값이므로, 서울서베이 개인 응답을 종속변수로 사용하는 HLM에 동시에 투입할 경우 공통 조사원천과 개념적 중복에서 오는 기계적 상관 및 내생성 우려가 있다. 따라서 주 분석에서는 상대적으로 외생적인 인프라·서비스 이용·경제적 부담 지표를 중심으로 모형을 구성하고, 제외된 세 차원은 UMC 지수 해석 및 민감도 분석에서 보조적으로 활용하였다.",
    )
    replace_startswith(
        doc,
        "분석은 4단계로 진행된다.",
        "분석은 4단계로 진행된다. Model 0은 무조건 평균모형 y_ij = beta_0j + e_ij, beta_0j = gamma_00 + u_0j로 급내상관계수(ICC)와 설계효과를 확인한다. Model 1은 개인 수준 공변량 X_ij를 고정효과로 추가한다. Model 2는 자치구 수준 UMC 변수 Z_j를 추가해 지역 조건의 평균적 연관을 추정한다. Model 3은 고령자(65세 이상)와 저학력 집단의 취약성을 검토하기 위해 X_ij x Z_j 교차수준 상호작용을 포함한다. 자치구별 절편은 무작위효과로 두고, 기울기는 고정효과로 추정하였다. 연령은 연속적 생애주기 기울기를 포착하기 위해 grand-mean centered 변수로 유지하고, elderly 더미는 65세 이상에서 나타날 수 있는 임계점형 취약성과 교차수준 상호작용을 식별하기 위해 별도로 포함하였다. Model 3에서는 Affordability 종합점수 대신 그 핵심 구성요소인 자치구 평균소득을 사용하여 경제적 접근성의 상호작용을 보다 직접적으로 검토하였다.",
    )
    replace_startswith(
        doc,
        "분석에 활용된 표본은 서울서베이",
        "분석에 활용된 표본은 서울서베이 2023-2024년 응답자 총 10,000명이며, 25개 자치구에서 평균적으로 400명이 표집되었다. 주요 변수에 대한 기술통계량은 <Table 3>에 제시하였다. 개인의 디지털 활용 점수 평균은 80점 만점에 60.88점으로 나타나 서울시민의 전반적 디지털 활용 수준이 높음을 보여준다. 종속변수는 음의 왜도가 일부 존재할 수 있으나, 기관 표준 척도와 10,000명 규모의 표본을 고려해 척도 선형화한 원점수대로 분석하였다.",
    )
    replace_startswith(
        doc,
        "그 외 개인 특성으로는, 여성이 남성 대비",
        "그 외 개인 특성으로는, 여성이 남성 대비 2.6점 높게 나타났다. 이는 국내 디지털 생활 서비스의 이용 패턴과 사회문화적 맥락에 따라 성별 차이가 달라질 수 있음을 시사한다. 장애인의 경우 3.1점 높게 나타났는데, 이는 표본 구성 및 서울시 장애인 지원 서비스 이용 경험과 관련된 선택 효과 가능성을 함께 고려해 해석해야 한다.",
    )
    replace_startswith(
        doc,
        "자치구 수준의 UMC 인프라 차원은 주효과만으로는",
        "자치구 수준의 UMC 인프라 차원은 주효과만으로는 대체로 유의하지 않았다. Model 2에서 세 차원 모두 통계적 유의성에 도달하지 못하였고, 우도비 검정 결과도 Model 1 대비 유의한 개선이 없었다. 이는 서울처럼 평균적 디지털 인프라 수준이 높은 도시에서는 지역 수준 변수가 전체 시민의 디지털 활용 평균을 크게 설명하지 못할 수 있음을 보여준다. 다만 교차수준 상호작용을 투입한 Model 3에서는 일부 지역 조건이 취약집단과 결합할 때 유의한 차이를 보였다.",
    )
    replace_startswith(
        doc,
        "본 분석의 핵심 발견은 교차수준 상호작용에서",
        "본 분석의 핵심 발견은 교차수준 상호작용에서 나타났다. 공공 WiFi 밀도와 고령 더미의 상호작용은 3.03점의 유의한 정적 효과를 보여, 공공 WiFi 인프라가 비고령자보다 고령자에게 상대적으로 더 큰 혜택을 제공할 가능성을 시사한다. Connectivity와 저학력 더미의 상호작용 역시 유의한 정적 효과를 보여, 네트워크 여건이 좋은 자치구에서 저학력 주민의 디지털 활용 격차가 일부 완화될 수 있음을 보여준다. 이는 인프라 투자가 평균 효과로만 나타나기보다 특정 취약집단에서 더 뚜렷한 연관을 가질 수 있음을 의미한다.",
    )
    replace_startswith(
        doc,
        "3.1절과 3.2절은 행정 통계와 설문 조사에 기반한",
        "3.1절과 3.2절은 행정 통계와 설문 조사에 기반한 정량 지표를 통해 서울의 디지털 격차를 측정하고 관련 요인을 분석하였다. 그러나 이러한 양적 접근만으로는 주민이 일상에서 디지털 연결성을 실제로 어떻게 경험하고 어떤 경로에서 공식 서비스 이용이 중단되는지 확인하기 어렵다. 3.3절은 비정형 텍스트를 분석 대상으로 삼아 이러한 표출 양상을 보완적으로 관찰하되, 플랫폼 게시글을 서울시민 전체의 대표 표본으로 간주하지 않는다.",
    )
    replace_startswith(
        doc,
        "이 절의 산출물은 세 가지다.",
        "이 절의 산출물은 세 가지다. 첫째, 개별 사례는 주민이 어떤 상황에서 어떤 연결성 문제를 경험했는지 보여준다. 둘째, 게시글의 이산 판정 결과를 자치구 수준으로 집계한 표출 프로파일은 텍스트가 지역별 UMC 패턴을 얼마나 반영하는지 검토할 수 있게 한다. 셋째, 추론 단계는 개별 사례가 왜 특정 지역 패턴의 일부로 읽힐 수 있는지를 설명하는 메커니즘 해석을 제공한다. 텍스트 분석 과정을 정리한 그림은 Figure 8과 같다.",
    )
    replace_startswith(
        doc,
        "분석의 텍스트 자료로는 한국의 지역 기반 커뮤니티 플랫폼",
        "분석의 텍스트 자료로는 한국의 지역 기반 커뮤니티 플랫폼인 당근의 동네생활 게시글을 활용하였다. 당근을 선택한 이유는 플랫폼 자체가 대표성을 갖기 때문이 아니라, 게시글에 행정동 및 자치구 수준의 지리적 메타데이터가 부여되어 3.1절과 3.2절의 공간 단위와 연결할 수 있기 때문이다. 따라서 본 자료는 디지털 결여의 전체 발생률이 아니라, 플랫폼 이용자가 지역 커뮤니티에 표출한 디지털 어려움의 구조를 관찰하는 자료로 해석한다.",
    )
    replace_startswith(
        doc,
        "원시 데이터는 서울 25개 자치구를 대상으로",
        "원시 데이터는 서울 25개 자치구를 대상으로 2024년 5월부터 2026년 3월까지 작성된 동네생활 게시글 중 수집 가능한 게시글을 대상으로 한다. 무작위 표본추출이 아니라, 기간·지역·게시판 범위를 사전에 정의한 뒤 접근 가능한 게시글을 포괄적으로 수집한 편의 표본에 가깝다. 중복 제거 후 총 1,287,761건이며, 삭제(DELETED) 663,163건(51.5%)과 차단(BLOCKED) 103,728건(8.1%) 등 본문 확인이 불가능한 게시글은 분석에서 제외하였다. 이는 결측치 대체의 대상이 아니라 플랫폼 접근 가능성에 따른 표본 탈락으로 처리하였다.",
    )
    replace_startswith(
        doc,
        "당근 데이터의 구조적 한계도 인식할 필요가 있다.",
        "당근 데이터의 구조적 한계도 명확하다. 플랫폼 이용자는 스마트폰과 지역 커뮤니티 앱을 사용할 수 있는 사람으로 한정되므로, 디지털 소외가 가장 심각한 집단은 과소 대표될 가능성이 높다. 또한 도움을 요청하거나 불편을 게시글로 표현하는 사람은 침묵하는 이용자와 다를 수 있어, 게시글은 '문제의 발생'이 아니라 '문제의 표출'을 측정한다. 본 연구에서는 개인을 식별할 수 있는 정보와 원문 노출을 최소화하고, 플랫폼의 이용 약관과 연구윤리 기준을 준수하는 범위에서 집계 및 익명화된 결과만 보고한다. 향후에는 플랫폼 이용자 인구구성과 서울시 전체 인구구성을 비교해 표본 편향의 크기를 정량적으로 점검할 필요가 있다.",
    )

    caption_replacements = {
        "[Figure 1] 분석 구조": "Figure 1. Regression-based analysis workflow",
        "[Figure 2] 프로젝트 개괄": "Figure 2. Project framework",
        "<Table 1> UMC Dimension Indicators": "Table 1. UMC Dimension Indicators, Units, Sources, and Construction",
        "Figure 4. UMC Dimension Scores": "Figure 4. UMC Dimension Scores",
        "Figure 6. Spatial autocorrelation by UMC level across Seoul’s 25 autonomous districts": "Figure 5. Spatial Autocorrelation by UMC Dimension across Seoul's 25 Districts",
        "<Table 2> Variable Summary": "Table 2. Variable Summary",
        "<Table 3> Descriptive Statistic": "Table 3. Descriptive Statistics",
        "<Table 4> HLM Estimation Results": "Table 4. HLM Estimation Results",
        "Figure #. 분류와 추론의 역할 구분": "Figure 6. Separation of Classification and Inference Roles",
        "표 #. 3.3절의 분석 단위와 해석 단위": "Table 5. Units of Analysis and Interpretation in Section 3.3",
        "Figure #. LLM 기반 텍스트 분석 파이프라인": "Figure 7. LLM-based Text Analysis Pipeline",
        "<Table 5> 가설 생성 결과: 결여 유형별 분포": "Table 7. Generated Hypotheses by Deficit Type",
    }
    for old, new in caption_replacements.items():
        replace_exact(doc, old, new)
    replace_startswith(doc, "텍스트 분류 결과는 Table #과 같다.", "텍스트 분류 결과는 Table 6과 같다.")
    replace_startswith(
        doc,
        "안전·보안(Safety & Security) 차원에서는 중랑구",
        "안전·보안(Safety & Security) 차원에서는 중랑구(+0.018)와 도봉구(+0.015)에서 피싱 사기 및 보이스피싱 관련 게시글의 비율이 높았다. 종로구는 실측 0.984로 최고 수준이면서 관련 게시글이 드물어 -0.019의 shift를 나타냈는데, 관광·상업 중심지의 주거 구성과 당근 이용자 집단의 특성이 상이할 가능성을 함께 고려할 필요가 있다. 6개 차원에 대한 사전-사후 비교 결과는 Figure 8과 같다.",
    )
    replace_startswith(
        doc,
        "본 절에서는 3.1절(UMC 지수 구축), 3.2절(HLM 다수준 분석), 3.3절",
        "본 절에서는 3.1절(UMC 지수 구축), 3.2절(HLM 다수준 분석), 3.3절(LLM 기반 텍스트 분석)의 결과를 종합하여, 서울의 디지털 격차가 어떤 구조로 작동하는지를 통합적으로 해석한다. 세 분석은 각각 '어디에 격차가 있는가'(측정), '무엇과 관련되는가'(회귀 기반 연관 분석), '격차가 어떻게 경험되는가'(질적 맥락)에 답하며, 이들의 교차 지점에서 정책적으로 유의미한 함의가 도출된다.",
    )
    replace_startswith(
        doc,
        "방법론적 관점에서, 세 분석의 결합은 각각의 한계를 상호 보완한다.",
        "방법론적 관점에서, 세 분석의 결합은 각각의 한계를 상호 보완한다. 3.1절의 종합지수는 자치구 간 상대적 위치를 파악하되 개인 내 변이를 포착하지 못하며, 3.2절의 HLM은 개인과 지역의 기여를 분리하되 디지털 격차의 질적 발현 경로를 설명하지 못한다. 3.3절의 텍스트 분석은 결여의 경험적 맥락을 가시화하되 대표성에 한계가 있다. 세 분석을 병렬이 아닌 순차적 삼각검증(측정 -> 회귀 기반 연관 분석 -> 질적 맥락)으로 설계함으로써, 하나의 자료원이 가진 약점을 다른 자료원이 보완하도록 하였다.",
    )


def prepare_image(src: Path, target_size: tuple[int, int], suffix: str) -> bytes:
    image = Image.open(src).convert("RGBA")
    tw, th = target_size
    scale = min(tw / image.width, th / image.height)
    new_size = (max(1, int(image.width * scale)), max(1, int(image.height * scale)))
    image = image.resize(new_size, Image.LANCZOS)
    canvas = Image.new("RGBA", (tw, th), (255, 255, 255, 255))
    canvas.alpha_composite(image, ((tw - new_size[0]) // 2, (th - new_size[1]) // 2))
    out = BytesIO()
    if suffix.lower() in {".jpg", ".jpeg"}:
        canvas.convert("RGB").save(out, format="JPEG", quality=95, subsampling=0)
    else:
        canvas.save(out, format="PNG")
    return out.getvalue()


def original_media_sizes(docx_path: Path) -> dict[str, tuple[int, int]]:
    sizes: dict[str, tuple[int, int]] = {}
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            with Image.open(BytesIO(zf.read(name))) as image:
                sizes[name] = image.size
    return sizes


def patch_package_media_and_xml(input_docx: Path, output_docx: Path) -> None:
    sizes = original_media_sizes(INPUT)
    replacements = {
        "word/media/image1.png": FIG_DIR / "fig_analysis_flow_en.png",
        "word/media/image10.png": FIG_DIR / "fig_context_seoul_map.png",
        "word/media/image12.png": REPORT_FIG_DIR / "fig_report_map_umc.png",
        "word/media/image13.png": FIG_DIR / "fig_prior_posterior_maps_en.png",
        "word/media/image15.png": FIG_DIR / "fig_connectivity_inputs_2024.png",
        "word/media/image16.png": FIG_DIR / "fig_normalization_formula_en.png",
        "word/media/image17.png": REPORT_FIG_DIR / "fig_report_bar_umc.png",
        "word/media/image7.png": REPORT_FIG_DIR / "fig_report_heatmap.png",
        "word/media/image8.jpg": FIG_DIR / "lisa_connectivity.png",
        "word/media/image2.jpg": FIG_DIR / "lisa_available_for_use.png",
        "word/media/image14.jpg": FIG_DIR / "lisa_affordability.png",
        "word/media/image3.jpg": FIG_DIR / "lisa_devices.png",
        "word/media/image5.jpg": FIG_DIR / "lisa_digital_skills.png",
        "word/media/image6.jpg": FIG_DIR / "lisa_safety.png",
        "word/media/image4.jpg": FIG_DIR / "fig_classification_inference_roles_en.png",
        "word/media/image9.png": FIG_DIR / "fig_bayesian_shift_maps_en.png",
    }
    with zipfile.ZipFile(input_docx, "r") as zin, zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in replacements and replacements[item.filename].exists():
                data = prepare_image(replacements[item.filename], sizes[item.filename], Path(item.filename).suffix)
            elif item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                text = re.sub(r'w:w="(-?[0-9]+)\.0"', r'w:w="\1"', text)
                text = re.sub(r'w:val="(-?[0-9]+)\.0"', r'w:val="\1"', text)
                text = re.sub(
                    r'w:val="([0-9]+)\.[0-9]+"',
                    lambda match: f'w:val="{round(float(match.group(1)))}"',
                    text,
                )
                data = text.encode("utf-8")
            zout.writestr(item, data)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(INPUT)
    replace_text(doc)
    rebuild_indicator_table(doc)
    update_variable_table(doc)
    polish_styles(doc)
    replace_problematic_tables_with_images(doc)
    resize_small_formula_images(doc)
    remove_trailing_empty_paragraphs(doc)
    doc.save(WORK)

    comments_strip = SKILL_DIR / "scripts/comments_strip.py"
    if comments_strip.exists():
        import subprocess

        subprocess.run(
            ["/Users/ujunbin/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3", str(comments_strip), str(WORK), "--out", str(STRIPPED)],
            check=True,
        )
        source = STRIPPED
    else:
        source = WORK
    patch_package_media_and_xml(source, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
