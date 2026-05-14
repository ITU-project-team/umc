"""KR 보고서 디자인 v2 — 전문 정책연구원 보고서 양식 업그레이드.

기존 v1 대비 강화 사항:
- 표지: 풀페이지 그라디언트 영역 + 메타정보 카드 + 액센트 라인
- 헤더/푸터: 듀얼톤 색상 띠 + 보고서 약칭 마크 + 페이지 N/M 형식
- 챕터 인트로: 큰 챕터 번호 글리프 + 섹션 키 메시지 박스
- 콜아웃 박스: '핵심 발견', '정의', '인용' 3종
- 빈 페이지: 챕터 종결 키 메시지 카드로 채우기
- 표: 강화된 헤더 라인 + 셀 패딩 개선
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from lxml.etree import _Element as _LxmlElement

INFOGRAPHICS_DIR = Path("manuscript/output/doc/infographics")

# ---- design constants ----
NAVY = "0E2A47"
NAVY_DARK = "0A1F35"
NAVY_LIGHT = "1E3A5F"
GOLD = "B8860B"
GOLD_LIGHT = "D4A82C"
TEAL = "2C7A7B"
TEAL_SOFT = "5BA3A4"
ACCENT_RED = "C53030"
TEXT_BLACK = "1A202C"
TEXT_GRAY = "4A5568"
TEXT_GRAY_LIGHT = "718096"
ZEBRA = "F4F6F8"
LIGHT_GRAY = "E2E8F0"
BG_CALLOUT_NAVY = "EDF2F7"
BG_CALLOUT_GOLD = "FFF8E6"
BG_CALLOUT_TEAL = "E6F4F4"
BG_CARD = "F8FAFC"

BODY_FONT = "KoPubWorld Dotum Medium"
BODY_FONT_FALLBACK = "KoPubWorld Dotum Medium"
HEAD_FONT = "KoPubWorld Dotum Bold"
HEAD_FONT_FALLBACK = "KoPubWorld Dotum Bold"

# 표 인덱스 분류 (renumber 적용 후)
FIGURE_CONTAINER_INDICES = {0, 2, 3}  # cover, fig pairs
SMALL_FONT_TABLE_INDICES = {1, 4, 5, 6, 7, 9, 10}

# 챕터 키 메시지 — 빈 페이지를 채울 카드 콘텐츠
CHAPTER_KEY_MESSAGES = {
    1: {
        "tag": "CHAPTER 1",
        "title": "프로젝트 개괄",
        "kicker": "서울 25개 자치구의 디지털 격차를",
        "body": "측정 → 설명 → 복원의 3단계 분석 체계로 검토한다. 비판적 실재론의 층화된 존재론을 메타이론으로 두어, 측정값·계수·텍스트가 동일한 디지털 연결성 현상의 다른 단면임을 명시한다.",
    },
    2: {
        "tag": "CHAPTER 2",
        "title": "이론적 기반과 해석적 맥락",
        "kicker": "사람 기반과 장소 기반 접근을 통합하는",
        "body": "장소민감적(place-sensitive) 접근의 학술·정책적 근거를 정리하고, 서울의 인구·인프라·정책 환경을 분석의 해석적 맥락으로 제시한다.",
    },
    3: {
        "tag": "CHAPTER 3",
        "title": "분석 결과",
        "kicker": "측정 → HLM → 베이지안 통합의",
        "body": "3단계 분석으로 서울 디지털 격차의 공간 분포, 개인-지역 기여 분리, 그리고 텍스트 기반 잠재 격차 신호까지 통합 추정한다.",
    },
    4: {
        "tag": "CHAPTER 4",
        "title": "정책 제언",
        "kicker": "관측·계수·텍스트의 3중 신호를 결합한",
        "body": "장소민감적 정책 처방을 제시하고, 본 프레임워크의 외적 타당성과 후속 연구를 위한 조건을 정리한다.",
    },
}


# ---- low-level helpers ----

def _set_raw_text(desc, value):
    try:
        _LxmlElement.text.fset(desc, value)
    except Exception:
        pass


def _scrub_text_attrs(elem) -> None:
    for desc in elem.iter():
        tlocal = desc.tag.split("}", 1)[-1] if isinstance(desc.tag, str) else ""
        if tlocal == "t":
            continue
        _set_raw_text(desc, None)


def _make(tag: str) -> OxmlElement:
    return OxmlElement(tag)


def _ensure_pPr(p_elem):
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = _make("w:pPr")
        p_elem.insert(0, pPr)
    return pPr


def _ensure_rPr_in_run(r_elem):
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = _make("w:rPr")
        r_elem.insert(0, rPr)
    return rPr


def _set_run_font(rPr, font_name: str, size_pt: float | None = None,
                  color_hex: str | None = None, bold: bool | None = None):
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _make("w:rFonts")
        rPr.insert(0, rFonts)
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(k), font_name)
    rFonts.set(qn("w:hint"), "eastAsia")
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = _make("w:lang")
        rPr.append(lang)
    lang.set(qn("w:eastAsia"), "ko-KR")
    if not lang.get(qn("w:val")):
        lang.set(qn("w:val"), "ko-KR")
    if size_pt is not None:
        sz = rPr.find(qn("w:sz"))
        if sz is None:
            sz = _make("w:sz")
            rPr.append(sz)
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = _make("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size_pt * 2)))
    if color_hex is not None:
        col = rPr.find(qn("w:color"))
        if col is None:
            col = _make("w:color")
            rPr.append(col)
        col.set(qn("w:val"), color_hex)
    if bold is True:
        if rPr.find(qn("w:b")) is None:
            rPr.append(_make("w:b"))
        if rPr.find(qn("w:bCs")) is None:
            rPr.append(_make("w:bCs"))
    elif bold is False:
        for tag in ("w:b", "w:bCs"):
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)


def _set_paragraph_spacing(pPr, line_240ths: int = 320,
                           before_pt: float | None = None,
                           after_pt: float | None = None,
                           line_rule: str = "auto"):
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = _make("w:spacing")
        pPr.append(spacing)
    if before_pt is not None:
        spacing.set(qn("w:before"), str(int(before_pt * 20)))
    if after_pt is not None:
        spacing.set(qn("w:after"), str(int(after_pt * 20)))
    spacing.set(qn("w:line"), str(line_240ths))
    spacing.set(qn("w:lineRule"), line_rule)


def _set_paragraph_alignment(pPr, align: str = "left"):
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = _make("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), align)


def _set_paragraph_shading(pPr, fill_hex: str):
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = _make("w:shd")
        pPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_paragraph_border(pPr, sides: dict):
    """sides = {'left': (sz_eighths, color_hex), ...}"""
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is not None:
        pPr.remove(pBdr)
    pBdr = _make("w:pBdr")
    pPr.append(pBdr)
    for side, (sz, color) in sides.items():
        b = _make(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)


def _set_paragraph_indent(pPr, left_pt: float | None = None,
                          right_pt: float | None = None,
                          first_line_pt: float | None = None):
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = _make("w:ind")
        pPr.append(ind)
    if left_pt is not None:
        ind.set(qn("w:left"), str(int(left_pt * 20)))
    if right_pt is not None:
        ind.set(qn("w:right"), str(int(right_pt * 20)))
    if first_line_pt is not None:
        ind.set(qn("w:firstLine"), str(int(first_line_pt * 20)))


def _make_run(text: str, font: str = BODY_FONT, size_pt: float = 10.5,
              color_hex: str = TEXT_BLACK, bold: bool = False) -> _LxmlElement:
    r = _make("w:r")
    rPr = _make("w:rPr")
    r.append(rPr)
    _set_run_font(rPr, font, size_pt=size_pt, color_hex=color_hex, bold=bold)
    t = _make("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_break_run(break_type: str = "page") -> _LxmlElement:
    r = _make("w:r")
    br = _make("w:br")
    br.set(qn("w:type"), break_type)
    r.append(br)
    return r


def _make_paragraph(runs: list, align: str = "left",
                    spacing_before: float = 0, spacing_after: float = 0,
                    line_240ths: int = 360,
                    shading: str | None = None,
                    border_sides: dict | None = None,
                    indent_left: float | None = None,
                    indent_right: float | None = None) -> _LxmlElement:
    p = _make("w:p")
    pPr = _make("w:pPr")
    p.append(pPr)
    _set_paragraph_alignment(pPr, align)
    _set_paragraph_spacing(pPr, line_240ths=line_240ths,
                           before_pt=spacing_before, after_pt=spacing_after)
    if shading:
        _set_paragraph_shading(pPr, shading)
    if border_sides:
        _set_paragraph_border(pPr, border_sides)
    if indent_left is not None or indent_right is not None:
        _set_paragraph_indent(pPr, left_pt=indent_left, right_pt=indent_right)
    for r in runs:
        p.append(r)
    return p


# ---- font enforcement ----

def force_run_fonts(doc) -> None:
    body = doc.element.body
    HEAD_STYLES = {"Heading1", "Heading2", "Heading3", "Heading4",
                   "1", "2", "3", "4",
                   "Heading 1", "Heading 2", "Heading 3", "Heading 4",
                   "Title", "Subtitle"}
    for p in body.iter(qn("w:p")):
        # skip our injected paragraphs (they are already styled)
        if p.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"):
            continue
        pPr = p.find(qn("w:pPr"))
        is_heading = False
        is_caption = False
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None:
                sid = ps.get(qn("w:val"))
                if sid in HEAD_STYLES:
                    is_heading = True
                if sid == "Caption":
                    is_caption = True
        if not is_heading and not is_caption:
            pPr_w = _ensure_pPr(p)
            sp = pPr_w.find(qn("w:spacing"))
            if sp is None:
                sp = _make("w:spacing")
                pPr_w.append(sp)
            # 1.3x line for KoPub Dotum (compact)
            sp.set(qn("w:line"), "300")
            sp.set(qn("w:lineRule"), "auto")
            sp.set(qn("w:after"), "20")
            sp.set(qn("w:before"), "0")
        for r in p.findall(qn("w:r")):
            rPr = _ensure_rPr_in_run(r)
            for rf in rPr.findall(qn("w:rFonts")):
                rPr.remove(rf)
            for tag in ("w:sz", "w:szCs"):
                for el in rPr.findall(qn(tag)):
                    rPr.remove(el)
            if not is_heading and not is_caption:
                for c in rPr.findall(qn("w:color")):
                    rPr.remove(c)
                for tag in ("w:spacing", "w:kern", "w:position"):
                    for el in rPr.findall(qn(tag)):
                        rPr.remove(el)
            if is_heading:
                _set_run_font(rPr, HEAD_FONT, size_pt=None,
                              color_hex=None, bold=None)
            elif is_caption:
                _set_run_font(rPr, HEAD_FONT, size_pt=10,
                              color_hex=NAVY, bold=True)
            else:
                _set_run_font(rPr, BODY_FONT, size_pt=10.5,
                              color_hex=TEXT_BLACK, bold=False)


def configure_styles(doc) -> None:
    for style in doc.styles:
        try:
            xml = style.element
        except AttributeError:
            continue
        tag = xml.tag.split("}", 1)[-1]
        if tag != "style":
            continue
        sid = xml.get(qn("w:styleId"))
        if sid is None:
            continue
        rPr = xml.find(qn("w:rPr"))
        pPr = xml.find(qn("w:pPr"))
        if sid in ("Normal", "Default", "DefaultParagraphFont"):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, BODY_FONT_FALLBACK, size_pt=10.5,
                          color_hex=TEXT_BLACK, bold=False)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=336, before_pt=0, after_pt=4)
        elif sid in ("Heading1", "1", "Heading 1"):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=24,
                          color_hex=NAVY, bold=True)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=320, before_pt=24, after_pt=14)
        elif sid in ("Heading2", "2", "Heading 2"):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=15,
                          color_hex=NAVY, bold=True)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=300, before_pt=20, after_pt=8)
        elif sid in ("Heading3", "3", "Heading 3"):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=12.5,
                          color_hex="2D3748", bold=True)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=300, before_pt=14, after_pt=4)
        elif sid in ("Heading4", "4", "Heading 4"):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=11.5,
                          color_hex=GOLD, bold=True)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=300, before_pt=10, after_pt=4)
        elif sid in ("Caption",):
            if rPr is None:
                rPr = _make("w:rPr"); xml.append(rPr)
            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=10,
                          color_hex=NAVY, bold=True)
            if pPr is None:
                pPr = _make("w:pPr"); xml.insert(0, pPr)
            _set_paragraph_spacing(pPr, line_240ths=300, before_pt=8, after_pt=6)
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = _make("w:jc")
                pPr.append(jc)
            jc.set(qn("w:val"), "center")
            keepNext = pPr.find(qn("w:keepNext"))
            if keepNext is None:
                keepNext = _make("w:keepNext")
                pPr.append(keepNext)


# ---- page setup ----

def _twips_from_mm(mm: float) -> int:
    return int(mm / 25.4 * 1440)


def configure_sections(doc) -> None:
    A4_W = _twips_from_mm(210)
    A4_H = _twips_from_mm(297)
    margins = dict(
        top=_twips_from_mm(28),
        bottom=_twips_from_mm(25),
        left=_twips_from_mm(25),
        right=_twips_from_mm(22),
        header=_twips_from_mm(15),
        footer=_twips_from_mm(15),
    )
    for sec in doc.sections:
        sectPr = sec._sectPr
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = _make("w:pgSz")
            sectPr.append(pgSz)
        pgSz.set(qn("w:w"), str(A4_W))
        pgSz.set(qn("w:h"), str(A4_H))
        if qn("w:orient") in pgSz.attrib:
            del pgSz.attrib[qn("w:orient")]
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is None:
            pgMar = _make("w:pgMar")
            sectPr.append(pgMar)
        pgMar.set(qn("w:top"), str(margins["top"]))
        pgMar.set(qn("w:right"), str(margins["right"]))
        pgMar.set(qn("w:bottom"), str(margins["bottom"]))
        pgMar.set(qn("w:left"), str(margins["left"]))
        pgMar.set(qn("w:header"), str(margins["header"]))
        pgMar.set(qn("w:footer"), str(margins["footer"]))


def _set_section_columns(sectPr, num: int = 1, space_twips: int = 453) -> None:
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = _make("w:cols")
        sectPr.append(cols)
    for child in list(cols):
        cols.remove(child)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def _remove_all_inline_section_breaks(body) -> None:
    to_remove = []
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        sp = pPr.find(qn("w:sectPr"))
        if sp is None:
            continue
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if text:
            continue
        to_remove.append(p)
    for p in to_remove:
        p.getparent().remove(p)


def apply_single_column_layout(doc) -> None:
    body = doc.element.body
    _remove_all_inline_section_breaks(body)
    for sec in doc.sections:
        _set_section_columns(sec._sectPr, num=1)


# ---- enhanced 3-line tables ----

def style_tables(doc) -> None:
    body = doc.element.body
    table_idx_map: dict = {}
    table_idx = 0
    for child in body.iterchildren():
        if child.tag.endswith("}tbl"):
            table_idx_map[id(child)] = table_idx
            table_idx += 1
    for tbl in [c for c in body.iterchildren() if c.tag.endswith("}tbl")]:
        rows = tbl.findall(qn("w:tr"))
        if not rows:
            continue
        idx = table_idx_map.get(id(tbl), -1)
        # 100% width
        tblPr_pre = tbl.find(qn("w:tblPr"))
        if tblPr_pre is None:
            tblPr_pre = _make("w:tblPr")
            tbl.insert(0, tblPr_pre)
        for tw in tblPr_pre.findall(qn("w:tblW")):
            tblPr_pre.remove(tw)
        tblW = _make("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr_pre.append(tblW)
        for tl in tblPr_pre.findall(qn("w:tblLayout")):
            tblPr_pre.remove(tl)
        layout = _make("w:tblLayout")
        layout.set(qn("w:type"), "autofit")
        tblPr_pre.append(layout)
        # cell margins (padding)
        for tcm in tblPr_pre.findall(qn("w:tblCellMar")):
            tblPr_pre.remove(tcm)
        tblCellMar = _make("w:tblCellMar")
        for side, val in (("top", 80), ("bottom", 80), ("left", 100), ("right", 100)):
            m = _make(f"w:{side}")
            m.set(qn("w:w"), str(val))
            m.set(qn("w:type"), "dxa")
            tblCellMar.append(m)
        tblPr_pre.append(tblCellMar)
        for tr in rows:
            for tc in tr.findall(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    continue
                for tcW in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(tcW)
        if idx in FIGURE_CONTAINER_INDICES:
            tblPr = tbl.find(qn("w:tblPr"))
            for tb in tblPr.findall(qn("w:tblBorders")):
                tblPr.remove(tb)
            tblBorders = _make("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = _make(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tblBorders.append(b)
            tblPr.append(tblBorders)
            for tr in rows:
                for tc in tr.findall(qn("w:tc")):
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = _make("w:tcPr")
                        tc.insert(0, tcPr)
                    for sh in tcPr.findall(qn("w:shd")):
                        tcPr.remove(sh)
                    for tb in tcPr.findall(qn("w:tcBorders")):
                        tcPr.remove(tb)
                    tcBorders = _make("w:tcBorders")
                    for side in ("top", "left", "bottom", "right"):
                        b = _make(f"w:{side}")
                        b.set(qn("w:val"), "nil")
                        tcBorders.append(b)
                    tcPr.append(tcBorders)
            continue
        # 3-line table with stronger top/bottom navy bars
        tblPr = tbl.find(qn("w:tblPr"))
        for tb in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(tb)
        tblBorders = _make("w:tblBorders")
        tblPr.append(tblBorders)
        # top: 1.5pt navy / bottom: 1pt navy / between header-and-row: navy 0.5pt
        for side, sz in (("top", 14), ("bottom", 10)):
            b = _make(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), NAVY)
            tblBorders.append(b)
        for side in ("left", "right", "insideV"):
            b = _make(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tblBorders.append(b)
        b = _make("w:insideH")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "CBD5E0")
        tblBorders.append(b)

        # zebra + header styling with stronger header bottom-border
        for ri, tr in enumerate(rows):
            cells = tr.findall(qn("w:tc"))
            for tc in cells:
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = _make("w:tcPr")
                    tc.insert(0, tcPr)
                for sh in tcPr.findall(qn("w:shd")):
                    tcPr.remove(sh)
                shd = _make("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                if ri == 0:
                    shd.set(qn("w:fill"), NAVY)
                elif ri % 2 == 1:
                    shd.set(qn("w:fill"), ZEBRA)
                else:
                    shd.set(qn("w:fill"), "FFFFFF")
                tcPr.append(shd)
                # vertical center for cleaner row alignment
                for vAlign in tcPr.findall(qn("w:vAlign")):
                    tcPr.remove(vAlign)
                vAlign = _make("w:vAlign")
                vAlign.set(qn("w:val"), "center")
                tcPr.append(vAlign)

                # header bottom border accent (gold underline below header)
                if ri == 0:
                    for tcb in tcPr.findall(qn("w:tcBorders")):
                        tcPr.remove(tcb)
                    tcBorders = _make("w:tcBorders")
                    bot = _make("w:bottom")
                    bot.set(qn("w:val"), "single")
                    bot.set(qn("w:sz"), "10")
                    bot.set(qn("w:color"), GOLD)
                    tcBorders.append(bot)
                    tcPr.append(tcBorders)

                cell_font_size = 9.5 if idx in SMALL_FONT_TABLE_INDICES else 10
                for p in tc.findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        rPr = _ensure_rPr_in_run(r)
                        if ri == 0:
                            _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=cell_font_size,
                                          color_hex="FFFFFF", bold=True)
                        else:
                            _set_run_font(rPr, BODY_FONT_FALLBACK, size_pt=cell_font_size,
                                          color_hex=TEXT_BLACK, bold=False)


# ---- header / footer ----

def add_headers_footers(doc) -> None:
    """전체 페이지 헤더/푸터 — 듀얼톤 띠 + 페이지 N/M + 보고서 약칭."""
    for si, sec in enumerate(doc.sections):
        sec.different_first_page_header_footer = True
        # ---- header ----
        header = sec.header
        hdr_elem = header._element
        for p in hdr_elem.findall(qn("w:p")):
            hdr_elem.remove(p)
        # 헤더는 두 줄: 윗줄=보고서 마크 + 챕터 식별, 아랫줄=썸 라인
        # 줄1: 좌측 보고서 약칭, 우측 챕터 식별
        new_p = _make("w:p")
        pPr = _make("w:pPr")
        new_p.append(pPr)
        # 탭 정렬: 좌측 + 우측
        tabs = _make("w:tabs")
        # 우측 탭 위치: A4 본문 폭 기준 (210mm - 25 - 22 = 163mm = 9242 twips)
        tab_right = _make("w:tab")
        tab_right.set(qn("w:val"), "right")
        tab_right.set(qn("w:pos"), "9242")
        tabs.append(tab_right)
        pPr.append(tabs)
        # 하단 굵은 navy bar + 가는 gold underline
        pBdr = _make("w:pBdr")
        bot = _make("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "12")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), NAVY)
        pBdr.append(bot)
        pPr.append(pBdr)

        # 좌측: 보고서 약칭 마크 (작은 박스 형태로)
        r1 = _make_run("ITU·UMC", font=HEAD_FONT_FALLBACK, size_pt=8.5,
                       color_hex=NAVY, bold=True)
        new_p.append(r1)
        r2 = _make_run("  ·  서울 디지털 격차 분석  ·  HIGH-FIVE",
                       font=HEAD_FONT_FALLBACK, size_pt=8.5,
                       color_hex=TEXT_GRAY_LIGHT, bold=False)
        new_p.append(r2)
        # 탭으로 우측 영역
        tab_run = _make("w:r")
        tab = _make("w:tab")
        tab_run.append(tab)
        new_p.append(tab_run)
        # 우측: 보고서 코드
        r3 = _make_run("KR · v1 · 2026.05",
                       font=HEAD_FONT_FALLBACK, size_pt=8.5,
                       color_hex=GOLD, bold=True)
        new_p.append(r3)
        hdr_elem.append(new_p)

        # ---- footer ----
        footer = sec.footer
        ftr_elem = footer._element
        for p in ftr_elem.findall(qn("w:p")):
            ftr_elem.remove(p)
        new_p = _make("w:p")
        pPr = _make("w:pPr")
        new_p.append(pPr)
        # 탭 양쪽 정렬
        tabs = _make("w:tabs")
        tab_center = _make("w:tab")
        tab_center.set(qn("w:val"), "center")
        tab_center.set(qn("w:pos"), "4621")
        tabs.append(tab_center)
        tab_right = _make("w:tab")
        tab_right.set(qn("w:val"), "right")
        tab_right.set(qn("w:pos"), "9242")
        tabs.append(tab_right)
        pPr.append(tabs)
        # 상단 0.5pt navy line
        pBdr = _make("w:pBdr")
        top = _make("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), NAVY)
        pBdr.append(top)
        pPr.append(pBdr)

        # 좌측: 보고서 영역 라벨
        r_left = _make_run("ITU UMC Data Hackathon 2026",
                           font=HEAD_FONT_FALLBACK, size_pt=8,
                           color_hex=TEXT_GRAY_LIGHT, bold=False)
        new_p.append(r_left)
        # 가운데로
        tr1 = _make("w:r")
        tr1.append(_make("w:tab"))
        new_p.append(tr1)
        # 가운데: PAGE N / NUMPAGES M
        for run_text, color_use in [("page_begin", None)]:
            pass
        # PAGE field
        r1 = _make("w:r")
        rPr1 = _make("w:rPr"); r1.append(rPr1)
        _set_run_font(rPr1, HEAD_FONT_FALLBACK, size_pt=10, color_hex=NAVY, bold=True)
        fldChar1 = _make("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        r1.append(fldChar1)
        new_p.append(r1)
        r2 = _make("w:r")
        rPr2 = _make("w:rPr"); r2.append(rPr2)
        _set_run_font(rPr2, HEAD_FONT_FALLBACK, size_pt=10, color_hex=NAVY, bold=True)
        instrText = _make("w:instrText")
        instrText.text = " PAGE   \\* MERGEFORMAT "
        instrText.set(qn("xml:space"), "preserve")
        r2.append(instrText)
        new_p.append(r2)
        r3 = _make("w:r")
        rPr3 = _make("w:rPr"); r3.append(rPr3)
        _set_run_font(rPr3, HEAD_FONT_FALLBACK, size_pt=10, color_hex=NAVY, bold=True)
        fldChar3 = _make("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        r3.append(fldChar3)
        new_p.append(r3)
        # 구분자
        r_sep = _make_run(" / ", font=HEAD_FONT_FALLBACK, size_pt=10,
                          color_hex=TEXT_GRAY_LIGHT, bold=False)
        new_p.append(r_sep)
        # NUMPAGES field
        r4 = _make("w:r")
        rPr4 = _make("w:rPr"); r4.append(rPr4)
        _set_run_font(rPr4, HEAD_FONT_FALLBACK, size_pt=9, color_hex=TEXT_GRAY, bold=False)
        fldChar4 = _make("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "begin")
        r4.append(fldChar4)
        new_p.append(r4)
        r5 = _make("w:r")
        rPr5 = _make("w:rPr"); r5.append(rPr5)
        _set_run_font(rPr5, HEAD_FONT_FALLBACK, size_pt=9, color_hex=TEXT_GRAY, bold=False)
        instrText2 = _make("w:instrText")
        instrText2.text = " NUMPAGES   \\* MERGEFORMAT "
        instrText2.set(qn("xml:space"), "preserve")
        r5.append(instrText2)
        new_p.append(r5)
        r6 = _make("w:r")
        rPr6 = _make("w:rPr"); r6.append(rPr6)
        _set_run_font(rPr6, HEAD_FONT_FALLBACK, size_pt=9, color_hex=TEXT_GRAY, bold=False)
        fldChar6 = _make("w:fldChar")
        fldChar6.set(qn("w:fldCharType"), "end")
        r6.append(fldChar6)
        new_p.append(r6)

        # 우측 탭
        tr2 = _make("w:r")
        tr2.append(_make("w:tab"))
        new_p.append(tr2)
        # 우측: 팀 마크
        r_right = _make_run("HIGH-FIVE",
                            font=HEAD_FONT_FALLBACK, size_pt=8.5,
                            color_hex=GOLD, bold=True)
        new_p.append(r_right)
        ftr_elem.append(new_p)

        # ---- first-page header/footer (cover): no header, special footer ----
        first_header = sec.first_page_header
        fh_elem = first_header._element
        for p in fh_elem.findall(qn("w:p")):
            fh_elem.remove(p)
        # 빈 paragraph (헤더 영역만 비움)
        empty = _make("w:p")
        fh_elem.append(empty)

        first_footer = sec.first_page_footer
        ff_elem = first_footer._element
        for p in ff_elem.findall(qn("w:p")):
            ff_elem.remove(p)
        empty = _make("w:p")
        ff_elem.append(empty)


# ---- cover upgrade ----

def upgrade_cover(doc) -> None:
    """표지 업그레이드: 풀 페이지 navy 영역 + 메타 카드 + 액센트 라인.

    추가로 표지 셀 직후의 메타 paragraphs (ITU UMC Data Hackathon..., Team..., March...)
    를 cell 외부에서 제거. 이들 메타는 cell 안에 새로 작성된 SEOUL · ... 박스로 대체된다.
    """
    body = doc.element.body
    tables = body.findall(qn("w:tbl"))
    if not tables:
        return
    cover = tables[0]
    # 표지 앞 빈 paragraphs 제거
    parent = cover.getparent()
    siblings_pre = list(parent.iterchildren())
    cover_idx_pre = siblings_pre.index(cover)
    for k in range(cover_idx_pre - 1, -1, -1):
        sib = siblings_pre[k]
        if not sib.tag.endswith("}p"):
            break
        text = "".join(t.text or "" for t in sib.iter(qn("w:t"))).strip()
        if text:
            break
        sib.getparent().remove(sib)
    # 표지 직후 외부 메타 paragraphs 제거
    siblings = list(parent.iterchildren())
    cover_idx = siblings.index(cover)
    META_KEYWORDS = ("ITU UMC Data Hackathon", "Team:", "Team :", "High-Five", "March 2026")
    to_remove = []
    j = cover_idx + 1
    while j < len(siblings):
        sib = siblings[j]
        if not sib.tag.endswith("}p"):
            break
        text = "".join(t.text or "" for t in sib.iter(qn("w:t"))).strip()
        if not text:
            # 빈 paragraph: 제거
            to_remove.append(sib)
            j += 1
            continue
        if any(kw in text for kw in META_KEYWORDS):
            to_remove.append(sib)
            j += 1
            continue
        break  # 첫 본문/헤딩 만나면 멈춤
    for el in to_remove:
        el.getparent().remove(el)
    rows = cover.findall(qn("w:tr"))
    if not rows:
        return
    # 행 높이 명시 (대형 표지 효과)
    for tr in rows:
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = _make("w:trPr")
            tr.insert(0, trPr)
        for trh in trPr.findall(qn("w:trHeight")):
            trPr.remove(trh)
        trh = _make("w:trHeight")
        # 풀 페이지 높이의 90% (A4 297mm × 0.85 ≒ 14000 twips)
        trh.set(qn("w:val"), "13800")
        trh.set(qn("w:hRule"), "exact")
        trPr.append(trh)
    cells = rows[0].findall(qn("w:tc"))
    if not cells:
        return
    tc = cells[0]
    # 셀 수직 가운데 정렬
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = _make("w:tcPr")
        tc.insert(0, tcPr)
    for vAlign in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(vAlign)
    vAlign = _make("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)
    # 대형 navy 영역
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = _make("w:tcPr")
        tc.insert(0, tcPr)
    for sh in tcPr.findall(qn("w:shd")):
        tcPr.remove(sh)
    shd = _make("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), NAVY)
    tcPr.append(shd)

    # 셀 높이를 키워 풀 영역 효과 (twips, 1cm = 567 twips)
    # tcMar: 큰 패딩
    for m in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(m)
    tcMar = _make("w:tcMar")
    for side, val in (("top", 1200), ("bottom", 1200), ("left", 600), ("right", 600)):
        em = _make(f"w:{side}")
        em.set(qn("w:w"), str(val))
        em.set(qn("w:type"), "dxa")
        tcMar.append(em)
    tcPr.append(tcMar)

    # 셀 안 paragraphs를 업그레이드
    paragraphs = tc.findall(qn("w:p"))
    # 첫 paragraph 앞에 GOLD 액센트 라벨 삽입
    gold_kicker = _make_paragraph(
        [_make_run("ITU UMC DATA HACKATHON 2026",
                   font=HEAD_FONT_FALLBACK, size_pt=10,
                   color_hex="D4A82C", bold=True)],
        align="center", spacing_before=0, spacing_after=8,
        line_240ths=240,
    )
    _mark_injected(gold_kicker)
    if paragraphs:
        paragraphs[0].addprevious(gold_kicker)

    for p in paragraphs:
        pPr = _ensure_pPr(p)
        _set_paragraph_alignment(pPr, "center")
        for r in p.findall(qn("w:r")):
            rPr = _ensure_rPr_in_run(r)
            text = "".join((t.text or "") for t in r.findall(qn("w:t")))
            if "Bridging" in text:
                _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=26,
                              color_hex="FFFFFF", bold=True)
            elif "Multilevel Analysis" in text or "across Seoul" in text:
                _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=13,
                              color_hex="D4A82C", bold=False)
            else:
                _set_run_font(rPr, HEAD_FONT_FALLBACK, size_pt=12,
                              color_hex="FFFFFF", bold=False)

    # 표지 셀 마지막에 골드 라인 + 메타 카드 + 디바이더 추가
    # 골드 디바이더
    divider_p = _make_paragraph(
        [_make_run("─────  ─────",
                   font=HEAD_FONT_FALLBACK, size_pt=12,
                   color_hex="D4A82C", bold=True)],
        align="center", spacing_before=20, spacing_after=8,
        line_240ths=240,
    )
    _mark_injected(divider_p)
    tc.append(divider_p)
    # 메타 카드
    meta_p = _make_paragraph(
        [
            _make_run("SEOUL · 25 DISTRICTS · MARCH 2026",
                      font=HEAD_FONT_FALLBACK, size_pt=10,
                      color_hex="FFFFFF", bold=True),
        ],
        align="center", spacing_before=0, spacing_after=4,
        line_240ths=240,
    )
    _mark_injected(meta_p)
    tc.append(meta_p)
    sub_p = _make_paragraph(
        [
            _make_run("REPORT v1 · KR EDITION",
                      font=HEAD_FONT_FALLBACK, size_pt=8.5,
                      color_hex="A0AEC0", bold=False),
        ],
        align="center", spacing_before=0, spacing_after=0,
        line_240ths=240,
    )
    _mark_injected(sub_p)
    tc.append(sub_p)

    # cell borders 제거
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = _make("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = _make(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # table borders 제거
    tblPr = cover.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = _make("w:tblPr")
        cover.insert(0, tblPr)
    for tb in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(tb)
    tblBorders = _make("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = _make(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---- chapter intro / key message cards ----

def _mark_injected(p_elem):
    p_elem.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected",
        "1",
    )


def _make_kicker_paragraph(text: str, color: str = GOLD) -> _LxmlElement:
    """챕터 식별 키커 (큰 NAVY/GOLD 라벨)."""
    runs = [_make_run(text, font=HEAD_FONT_FALLBACK, size_pt=10,
                      color_hex=color, bold=True)]
    p = _make_paragraph(
        runs, align="left", spacing_before=0, spacing_after=4,
        line_240ths=240,
    )
    _mark_injected(p)
    return p


def _make_chapter_card(meta: dict) -> list:
    """챕터 인트로 카드: 큰 챕터 번호 + 키커 + 본문 (정렬 통일)."""
    paragraphs = []
    # 1. 키커 (CHAPTER N) — 좌측 정렬, 카드와 동일 좌측 들여쓰기 0
    p_kicker = _make_paragraph(
        [_make_run(meta["tag"], font=HEAD_FONT_FALLBACK, size_pt=10,
                   color_hex=GOLD, bold=True)],
        align="left", spacing_before=0, spacing_after=4,
        line_240ths=240,
    )
    _mark_injected(p_kicker)
    paragraphs.append(p_kicker)
    # 2. 메인 키 메시지 박스 (좌측 4pt navy bar) — 들여쓰기 12pt 통일
    runs_kicker = [
        _make_run(meta["kicker"], font=HEAD_FONT_FALLBACK,
                  size_pt=12.5, color_hex=NAVY, bold=True),
    ]
    p1 = _make_paragraph(
        runs_kicker, align="left", spacing_before=6, spacing_after=4,
        line_240ths=320,
        shading=BG_CALLOUT_NAVY,
        border_sides={"left": (24, NAVY), "top": (4, NAVY_LIGHT),
                      "right": (4, NAVY_LIGHT), "bottom": (0, NAVY_LIGHT)},
        indent_left=12, indent_right=12,
    )
    _mark_injected(p1)
    paragraphs.append(p1)
    # 3. 본문 박스 (이어지는 영역) — 동일 들여쓰기 12pt
    runs_body = [
        _make_run(meta["body"], font=BODY_FONT_FALLBACK,
                  size_pt=10.5, color_hex=TEXT_BLACK, bold=False),
    ]
    p2 = _make_paragraph(
        runs_body, align="justify", spacing_before=0, spacing_after=14,
        line_240ths=320,
        shading=BG_CALLOUT_NAVY,
        border_sides={"left": (24, NAVY), "top": (0, NAVY_LIGHT),
                      "right": (4, NAVY_LIGHT), "bottom": (4, NAVY_LIGHT)},
        indent_left=12, indent_right=12,
    )
    _mark_injected(p2)
    paragraphs.append(p2)
    # 4. 한 줄 여백
    p3 = _make_paragraph([], spacing_before=0, spacing_after=8, line_240ths=240)
    _mark_injected(p3)
    paragraphs.append(p3)
    return paragraphs


def inject_chapter_intros(doc) -> None:
    """Heading1 직후 챕터 인트로 카드 삽입.

    Heading1 텍스트 패턴: '1. 개요', '2. 이론적 기반과 해석적 맥락', '3. 분석 결과' 등
    Heading1 직전에 페이지 브레이크를 강제하여 새 페이지에서 챕터 시작.
    """
    body = doc.element.body
    inserts = []
    for p in body.iter(qn("w:p")):
        if p.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"):
            continue
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        ps = pPr.find(qn("w:pStyle"))
        if ps is None:
            continue
        if ps.get(qn("w:val")) not in ("1", "Heading1", "Heading 1"):
            continue
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        chap_num = None
        if text and text[0].isdigit():
            try:
                chap_num = int(text.split(".")[0])
            except ValueError:
                pass
        if chap_num is None or chap_num not in CHAPTER_KEY_MESSAGES:
            continue
        meta = CHAPTER_KEY_MESSAGES[chap_num]
        cards = _make_chapter_card(meta)
        inserts.append((p, cards, chap_num))

    for anchor, paragraphs, chap_num in inserts:
        parent = anchor.getparent()
        idx = list(parent).index(anchor)
        # Heading1 자체에 pageBreakBefore 추가 (챕터 1은 표지 직후라 의도적 예외 가능)
        pPr = anchor.find(qn("w:pPr"))
        if pPr is None:
            pPr = _make("w:pPr")
            anchor.insert(0, pPr)
        if pPr.find(qn("w:pageBreakBefore")) is None:
            pbb = _make("w:pageBreakBefore")
            # pPr 시작 부근에 삽입
            pPr.insert(0, pbb)
        # cards 삽입 (heading 다음 위치)
        for offset, new_p in enumerate(paragraphs, start=1):
            parent.insert(idx + offset, new_p)


# ---- callout boxes for key text patterns ----

def _make_callout_box(label: str, body_text: str,
                      label_color: str = NAVY,
                      bg: str = BG_CALLOUT_NAVY,
                      bar_color: str = NAVY) -> list:
    """콜아웃 박스 paragraph 리스트 (정렬 통일)."""
    paragraphs = []
    # 라벨 — 좌측 12pt 들여쓰기 (챕터 카드와 일치)
    runs_label = [_make_run(label, font=HEAD_FONT_FALLBACK,
                            size_pt=10, color_hex=label_color, bold=True)]
    p1 = _make_paragraph(
        runs_label, align="left", spacing_before=8, spacing_after=2,
        line_240ths=260,
        shading=bg,
        border_sides={"left": (24, bar_color), "top": (4, bar_color),
                      "right": (4, bar_color), "bottom": (0, bar_color)},
        indent_left=12, indent_right=12,
    )
    _mark_injected(p1)
    paragraphs.append(p1)
    # 본문 — 동일 들여쓰기 12pt
    runs_body = [_make_run(body_text, font=BODY_FONT_FALLBACK,
                           size_pt=10, color_hex=TEXT_BLACK, bold=False)]
    p2 = _make_paragraph(
        runs_body, align="justify", spacing_before=0, spacing_after=10,
        line_240ths=300,
        shading=bg,
        border_sides={"left": (24, bar_color), "top": (0, bar_color),
                      "right": (4, bar_color), "bottom": (4, bar_color)},
        indent_left=12, indent_right=12,
    )
    _mark_injected(p2)
    paragraphs.append(p2)
    return paragraphs


# ---- end-of-chapter spacer cards ----

def _make_end_of_chapter_card(chap_num: int) -> list:
    """장 종결 후 빈 페이지를 채우는 키 메시지 카드."""
    if chap_num not in CHAPTER_KEY_MESSAGES:
        return []
    meta = CHAPTER_KEY_MESSAGES[chap_num]
    paragraphs = []
    # 한 줄 여백
    paragraphs.append(_mark_injected_p(_make_paragraph([], spacing_before=12, spacing_after=0)))
    # 작은 라벨
    runs_label = [_make_run("END OF " + meta["tag"], font=HEAD_FONT_FALLBACK,
                            size_pt=9, color_hex=GOLD, bold=True)]
    p_lbl = _make_paragraph(
        runs_label, align="center", spacing_before=0, spacing_after=4,
        line_240ths=240,
    )
    _mark_injected(p_lbl)
    paragraphs.append(p_lbl)
    # 큰 키 메시지
    runs_msg = [_make_run("· " + meta["title"] + " ·",
                          font=HEAD_FONT_FALLBACK, size_pt=18,
                          color_hex=NAVY, bold=True)]
    p_msg = _make_paragraph(
        runs_msg, align="center", spacing_before=0, spacing_after=10,
        line_240ths=300,
    )
    _mark_injected(p_msg)
    paragraphs.append(p_msg)
    # 짧은 한 줄
    runs_sub = [_make_run(meta["kicker"], font=BODY_FONT_FALLBACK,
                          size_pt=11, color_hex=TEXT_GRAY, bold=False)]
    p_sub = _make_paragraph(
        runs_sub, align="center", spacing_before=0, spacing_after=20,
        line_240ths=300,
    )
    _mark_injected(p_sub)
    paragraphs.append(p_sub)
    return paragraphs


def _mark_injected_p(p):
    _mark_injected(p)
    return p


# ---- image insertion via python-docx ----

def _make_image_paragraph(doc, image_path: Path, width_cm: float = 16.0,
                          caption: str | None = None) -> list:
    """이미지를 paragraph로 삽입 (캡션 옵션) — 가운데 정렬, 캡션 keepNext."""
    paragraphs = []
    p = doc.add_paragraph()
    p.alignment = 1  # CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    img_p = p._p
    body = doc.element.body
    body.remove(img_p)
    pPr = _ensure_pPr(img_p)
    _set_paragraph_alignment(pPr, "center")
    _set_paragraph_spacing(pPr, line_240ths=240, before_pt=10, after_pt=4)
    # keepNext so caption stays on same page as image
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(_make("w:keepNext"))
    _mark_injected(img_p)
    paragraphs.append(img_p)
    if caption:
        cap = _make_paragraph(
            [_make_run(caption, font=HEAD_FONT_FALLBACK, size_pt=9.5,
                       color_hex=NAVY, bold=True)],
            align="center", spacing_before=2, spacing_after=14,
            line_240ths=240,
        )
        _mark_injected(cap)
        paragraphs.append(cap)
    return paragraphs


def inject_infographics(doc) -> None:
    """본문 핵심 위치에 인포그래픽 자산을 삽입.

    삽입 위치:
    - '1.1 이론적 이해' 헤딩 직전 (UMC 차원 인포그래픽 + 분석 파이프라인)
    - '3. 분석 결과' 카드 다음 (key findings card)
    - '4. 정책 제언' 카드 다음 (policy summary)
    """
    body = doc.element.body

    inserts = []  # (anchor_paragraph, [paragraphs_to_insert_AFTER])

    # 인포그래픽 후보 위치 찾기
    for p in body.iter(qn("w:p")):
        if p.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"):
            continue
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        ps = pPr.find(qn("w:pStyle"))
        if ps is None:
            continue
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        sid = ps.get(qn("w:val"))
        # 1.2 분석 방법 이해 다음에 분석 파이프라인 인포그래픽
        # → 헤딩 자체 이전이 아니라, 1.2 헤딩 직전에 삽입할 수도 있음
        # 안전하게는, 1장 챕터 카드 직후에 인포그래픽 삽입
        # 하지만 위에서는 anchor_paragraph + AFTER 방식이므로
        # 1장 H1 paragraph anchor 후 추가
        pass

    # Heading1 (1, 2, 3, 4) 위치 찾기
    h1_paragraphs = {}  # chap_num -> p
    for p in body.iter(qn("w:p")):
        if p.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"):
            continue
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        ps = pPr.find(qn("w:pStyle"))
        if ps is None:
            continue
        sid = ps.get(qn("w:val"))
        if sid not in ("1", "Heading1", "Heading 1"):
            continue
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if text and text[0].isdigit():
            try:
                cn = int(text.split(".")[0])
                h1_paragraphs[cn] = p
            except ValueError:
                pass

    # 챕터 1 H1 anchor에서 — chapter card 다음 위치에 UMC + Pipeline 인포그래픽
    if 1 in h1_paragraphs:
        umc_path = INFOGRAPHICS_DIR / "umc_dimensions_hex.png"
        pipe_path = INFOGRAPHICS_DIR / "analysis_pipeline.png"
        if umc_path.exists() and pipe_path.exists():
            paras = []
            paras.extend(_make_image_paragraph(
                doc, umc_path, width_cm=15.5,
                caption="그림 0-1. UMC 6개 차원 측정 체계",
            ))
            paras.extend(_make_image_paragraph(
                doc, pipe_path, width_cm=16.0,
                caption="그림 0-2. 3단계 분석 파이프라인",
            ))
            inserts.append((h1_paragraphs[1], paras, "chapter_card_end"))

    # 챕터 3 H1 anchor에서 — chapter card 다음 위치에 Key findings card
    if 3 in h1_paragraphs:
        kf_path = INFOGRAPHICS_DIR / "key_findings_card.png"
        dd_path = INFOGRAPHICS_DIR / "digital_desert_summary.png"
        if kf_path.exists():
            paras = []
            paras.extend(_make_image_paragraph(
                doc, kf_path, width_cm=16.0,
                caption="그림 0-3. 핵심 발견 요약",
            ))
            if dd_path.exists():
                paras.extend(_make_image_paragraph(
                    doc, dd_path, width_cm=16.0,
                    caption="그림 0-4. 디지털 사막 — 종합지수 하위 4개 자치구",
                ))
            inserts.append((h1_paragraphs[3], paras, "chapter_card_end"))

    # 챕터 4 H1 anchor에서 — chapter card 다음 위치에 Policy summary
    if 4 in h1_paragraphs:
        ps_path = INFOGRAPHICS_DIR / "chapter4_policy_summary.png"
        if ps_path.exists():
            paras = []
            paras.extend(_make_image_paragraph(
                doc, ps_path, width_cm=16.0,
                caption="그림 0-5. 장소민감적 정책 제언 4대 축",
            ))
            inserts.append((h1_paragraphs[4], paras, "chapter_card_end"))

    # 삽입: 챕터 카드(이미 삽입된 paragraphs 4개) 다음에 추가
    # inject_chapter_intros가 anchor 다음 4개 paragraph를 삽입했으므로,
    # 우리는 anchor + 4 위치 다음에 삽입
    for anchor, paragraphs_to_insert, mode in inserts:
        parent = anchor.getparent()
        siblings = list(parent)
        anchor_idx = siblings.index(anchor)
        # chapter card는 anchor 다음 4개 paragraphs (kicker, navy box-header, body, spacer)
        # spacer 다음 위치에 인포그래픽
        offset = 4
        for j, np in enumerate(paragraphs_to_insert):
            parent.insert(anchor_idx + 1 + offset + j, np)


# ---- callout box injection ----

CALLOUT_TRIGGERS = [
    # (paragraph keyword set [needs to be inside one paragraph], label, ...)
    {
        "anchor_keyword": "잠정적으로 디지털 사막",
        "label": "핵심 발견",
        "color_label": NAVY,
        "bg": BG_CALLOUT_NAVY,
        "bar": NAVY,
        "title": "디지털 사막 4개 자치구",
        "body": "종합지수 하위 자치구(중랑·도봉·강북·노원)는 Devices·Safety 차원에서 동반 결핍을 보이며, 공간적으로도 서울 북부에 집중되어 있다. 이는 권역 단위 협력 거버넌스 구축의 정당화 근거를 제공한다.",
    },
    {
        "anchor_keyword": "Bhaskar, 1975",
        "label": "메타이론",
        "color_label": TEAL,
        "bg": BG_CALLOUT_TEAL,
        "bar": TEAL,
        "title": "층화된 존재론 — 경험·실재·실제 영역",
        "body": "관측된 디지털 격차(경험 영역)는 사회·경제·물리적 메커니즘(실재 영역)이 특정 조건에서 작동한 결과(실제 영역)이다. 본 분석은 세 영역을 명시적으로 분리하여 측정·계수·텍스트 산출물의 인식론적 위치를 다르게 둔다.",
    },
    {
        "anchor_keyword": "통합적 정책",
        "label": "정책 시사점",
        "color_label": GOLD,
        "bg": BG_CALLOUT_GOLD,
        "bar": GOLD,
        "title": "장소민감적 처방의 4대 축",
        "body": "디지털 사막 우선 개입 → 차원별 차등 처방 → 사람·장소 통합 접근 → 정량+텍스트 모니터링. 이 4대 축은 3장의 측정·HLM·텍스트 결과를 하나의 정책 프레임으로 통합한다.",
    },
]


def inject_callouts(doc) -> None:
    """본문 핵심 위치에 콜아웃 박스 삽입.

    트리거 키워드를 포함한 paragraph를 찾아 그 직전에 callout box paragraphs 삽입.
    """
    body = doc.element.body
    inserts = []
    for trigger in CALLOUT_TRIGGERS:
        kw = trigger["anchor_keyword"]
        for p in body.iter(qn("w:p")):
            if p.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"):
                continue
            text = "".join(t.text or "" for t in p.iter(qn("w:t")))
            if kw in text:
                # 이미 처리된 anchor면 skip
                if any(anchor is p for anchor, _ in inserts):
                    break
                callout = _make_callout_box(
                    label=trigger["label"] + " · " + trigger["title"],
                    body_text=trigger["body"],
                    label_color=trigger["color_label"],
                    bg=trigger["bg"],
                    bar_color=trigger["bar"],
                )
                inserts.append((p, callout))
                break
    for anchor, paragraphs in inserts:
        parent = anchor.getparent()
        idx = list(parent).index(anchor)
        for offset, np in enumerate(paragraphs):
            parent.insert(idx + offset, np)


# ---- main pipeline ----

def main(argv: list[str]) -> int:
    if len(argv) < 3:
        print("usage: kr_report_redesign_v2.py <input.docx> <output.docx>")
        return 2
    src = Path(argv[1])
    dst = Path(argv[2])
    dst.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(str(src))
    # 글로벌 raw text scrub
    for p in doc.element.body.iter(qn("w:p")):
        _scrub_text_attrs(p)

    configure_styles(doc)
    force_run_fonts(doc)
    configure_sections(doc)
    apply_single_column_layout(doc)
    style_tables(doc)
    add_headers_footers(doc)
    upgrade_cover(doc)
    inject_chapter_intros(doc)
    inject_infographics(doc)
    inject_callouts(doc)

    # remove our internal marker attributes before save (Word strict mode)
    UMC_MARKER = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcInjected"
    )
    UMC_WRAP = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}umcWrap"
    )
    for el in doc.element.body.iter():
        if UMC_MARKER in el.attrib:
            del el.attrib[UMC_MARKER]
        if UMC_WRAP in el.attrib:
            del el.attrib[UMC_WRAP]

    # OOXML strict: enforce w:pPr child order so Word can open the file
    enforce_ppr_child_order(doc)
    # ensure every wp:docPr / pic:cNvPr has unique id (Word rejects duplicates)
    deduplicate_drawing_ids(doc)
    # fix inline pictures so Word strict mode accepts them
    fix_inline_pictures(doc)
    # remove embedded fonts (0-byte font parts trigger Word "file corrupt")
    strip_embedded_fonts(doc)

    doc.save(str(dst))
    # post-save: remove font/odttf parts and fontTable rels from the docx zip
    purge_embedded_font_parts(dst)
    # LibreOffice round-trip cleans up any remaining OOXML quirks Word rejects
    libreoffice_roundtrip(dst)
    print(f"saved: {dst}")
    return 0


def libreoffice_roundtrip(docx_path) -> None:
    """Re-save via LibreOffice's MS Word 2007 XML filter to normalize OOXML."""
    import subprocess
    import shutil
    import tempfile
    from pathlib import Path
    p = Path(str(docx_path))
    with tempfile.TemporaryDirectory() as td:
        try:
            subprocess.run(
                [
                    "/opt/homebrew/bin/soffice",
                    "--headless",
                    "--convert-to", "docx:MS Word 2007 XML",
                    str(p),
                    "--outdir", td,
                ],
                check=True, capture_output=True, timeout=180,
            )
        except Exception as e:
            print(f"  (warning) libreoffice roundtrip skipped: {e}")
            return
        out = Path(td) / p.name
        if out.exists() and out.stat().st_size > 0:
            shutil.copy(str(out), str(p))
    # purge any embedded fonts LibreOffice may have re-added
    purge_embedded_font_parts(p)


def deduplicate_drawing_ids(doc) -> None:
    """Ensure every <wp:docPr id="..."/> and <pic:cNvPr id="..."/> has a unique id.

    Word strict mode rejects documents with duplicate drawing/picture ids.
    """
    body = doc.element.body
    WP_DOCPR = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr"
    PIC_CNVPR = "{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr"
    next_id = [1000]
    seen = set()
    for tag in (WP_DOCPR, PIC_CNVPR):
        for el in body.iter(tag):
            cur = el.get("id")
            if cur is None or cur in seen or cur == "0":
                while str(next_id[0]) in seen:
                    next_id[0] += 1
                el.set("id", str(next_id[0]))
                seen.add(str(next_id[0]))
                next_id[0] += 1
            else:
                seen.add(cur)


def fix_inline_pictures(doc) -> None:
    """python-docx의 add_picture가 생성한 inline drawing을 Word-strict 호환 형식으로 보강.

    누락 가능 요소:
    - wp:inline: distT/B/L/R="0" 속성
    - wp:effectExtent
    - a:prstGeom 안 a:avLst
    """
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    body = doc.element.body
    for inline in body.iter("{%s}inline" % WP):
        # add dist* attributes
        for k in ("distT", "distB", "distL", "distR"):
            if k not in inline.attrib:
                inline.set(k, "0")
        # ensure wp:effectExtent right after wp:extent
        extent = inline.find("{%s}extent" % WP)
        eff = inline.find("{%s}effectExtent" % WP)
        if extent is not None and eff is None:
            new_eff = OxmlElement("wp:effectExtent")
            new_eff.set("l", "0")
            new_eff.set("t", "0")
            new_eff.set("r", "0")
            new_eff.set("b", "0")
            extent.addnext(new_eff)
    # a:prstGeom must contain a:avLst (Word strict requires it)
    for prst in body.iter("{%s}prstGeom" % A):
        avlst = prst.find("{%s}avLst" % A)
        if avlst is None:
            new_avlst = OxmlElement("a:avLst")
            prst.append(new_avlst)


def strip_embedded_fonts(doc) -> None:
    """fontTable.xml에서 모든 임베디드 폰트 참조 제거."""
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except Exception:
        RT = None
    try:
        ft = doc.part.related_parts
    except Exception:
        ft = None
    # fontTable.xml 직접 편집
    for rel_id, rel in list(doc.part.rels.items()):
        target = getattr(rel, "target_ref", "") or getattr(getattr(rel, "_target", None), "partname", "")
        target_str = str(target)
        if "fontTable" in target_str:
            ft_part = rel.target_part
            from lxml import etree
            tree = etree.fromstring(ft_part.blob)
            W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            # remove embedRegular/Bold/Italic/BoldItalic from each w:font
            for fnt in tree.findall(W + "font"):
                for tag in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
                    for emb in fnt.findall(W + tag):
                        fnt.remove(emb)
            new_blob = etree.tostring(tree, xml_declaration=True, encoding="UTF-8",
                                      standalone=True)
            ft_part._blob = new_blob


def purge_embedded_font_parts(docx_path) -> None:
    """저장된 .docx zip에서 fonts/*.odttf 파트와 fontTable.xml.rels의 폰트 관계 제거.

    Word는 fontTable.xml.rels에 명시된 r:id의 .odttf 파일이 0바이트면 손상으로 인식.
    가장 안전한 해결: 임베디드 폰트 파트를 zip에서 완전히 제거하고 rels도 정리.
    """
    import zipfile
    import shutil
    from pathlib import Path
    p = Path(str(docx_path))
    tmp = p.with_suffix(".docx.tmp")
    with zipfile.ZipFile(p, "r") as zin:
        names = zin.namelist()
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                # skip embedded font binary parts
                if n.startswith("word/fonts/") and n.endswith(".odttf"):
                    continue
                data = zin.read(n)
                # rewrite fontTable.xml.rels: remove font relationships
                if n == "word/_rels/fontTable.xml.rels":
                    text = data.decode("utf-8")
                    import re
                    # remove every <Relationship .../> whose Type contains '/font'
                    text = re.sub(
                        r'<Relationship[^>]*?Type="[^"]*/font"[^>]*?/>',
                        "", text,
                    )
                    data = text.encode("utf-8")
                elif n == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    import re
                    # remove every <Override .../> for /word/fonts/*.odttf
                    text = re.sub(
                        r'<Override\s+PartName="/word/fonts/[^"]+\.odttf"[^>]*?/>',
                        "", text,
                    )
                    data = text.encode("utf-8")
                zout.writestr(n, data)
    shutil.move(str(tmp), str(p))


# OOXML CT_PPrBase child order — Word rejects out-of-order children
_PPR_CHILD_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
    "pPrChange",
]


def enforce_ppr_child_order(doc) -> None:
    """Sort w:pPr children into the spec-required order. Word strict mode rejects
    paragraphs whose pPr children are out of CT_PPrBase order."""
    body = doc.element.body
    order_map = {name: i for i, name in enumerate(_PPR_CHILD_ORDER)}
    # iterate every pPr in body, headers, footers
    parts = [body]
    # also include section header/footer parts
    for sec in doc.sections:
        try:
            parts.append(sec.header._element)
            parts.append(sec.footer._element)
            parts.append(sec.first_page_header._element)
            parts.append(sec.first_page_footer._element)
        except Exception:
            pass
    for root in parts:
        for pPr in root.iter(qn("w:pPr")):
            children = list(pPr)
            def keyfn(el):
                tlocal = el.tag.split("}", 1)[-1]
                return order_map.get(tlocal, 999)
            sorted_children = sorted(children, key=keyfn)
            if sorted_children != children:
                for c in children:
                    pPr.remove(c)
                for c in sorted_children:
                    pPr.append(c)


if __name__ == "__main__":
    sys.exit(main(sys.argv))
